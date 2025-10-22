import os
import PyPDF2
import docx
import re


class Readfile():
    """
    Класс получает файл, проверяет его тип. Если тип подходит, файл открывается и текст из него сохраняется в БД (будет осхраняться)
    """

    def __init__(self, name):
        self.name = name
        self.allowed_extensions = {'.txt', '.pdf', '.doc', '.docx'}
        self.content = None

    def _check_file_type(self):
        _, extension = os.path.splitext(self.name)
        if extension.lower() not in self.allowed_extensions:
            raise ValueError(
                f"Неподдерживаемый тип файла: {extension}. Разрешены: {', '.join(self.allowed_extensions)}")

        if not os.path.exists(self.name):
            raise FileNotFoundError(f"Файл {self.name} не найден")

        return extension.lower()

    def clear(self):

        patterns = [
            r'Рис\.\s*\d+\.\s*[^\n]+',  # Рис. 1. Описание
            r'Рисунок\s*\d+\.\s*[^\n]+',  # Рисунок 1. Описание
            r'Fig\.\s*\d+\.\s*[^\n]+',  # Fig. 1. Description
            r'Figure\s*\d+\.\s*[^\n]+',  # Figure 1. Description
            r'РИС\.\s*\d+\.\s*[^\n]+',  # РИС. 1. ОПИСАНИЕ
        ]

        cleaned_text = self.content
        for pattern in patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE)

        # Очистка результата
        cleaned_text = re.sub(r'\n\s*\n', '\n', cleaned_text)
        cleaned_text = re.sub(r' +', ' ', cleaned_text)
        cleaned_text = cleaned_text.strip()

        self.content = cleaned_text

    def read(self):
        try:
            extension = self._check_file_type()

            if extension == '.txt':
                self._read_txt()
            elif extension == '.pdf':
                self._read_pdf()
            elif extension in ['.doc', '.docx']:
                self._read_docx()



        except Exception as e:
            raise Exception(f"Ошибка при чтении файла: {str(e)}")

    def _read_txt(self):
        with open(self.name, 'r', encoding='utf-8') as file:
            self.content = file.read()

    def _read_pdf(self):
        try:
            with open(self.name, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                self.content = text
        except Exception as e:
            raise Exception(f"Ошибка чтения PDF: {str(e)}")

    def _read_docx(self):
        """Читает DOC/DOCX файл"""
        try:
            doc = docx.Document(self.name)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
            for table in doc.tables:
                text += self._read_docx_table(table)
            self.content = text
        except Exception as e:
            raise Exception(f"Ошибка чтения DOCX: {str(e)}")

    def _read_docx_table(self, table):
        """Извлекает текст из таблицы DOCX"""
        table_text = ""

        for row in table.rows:
            row_text = []
            for cell in row.cells:
                # Рекурсивно обрабатываем вложенные таблицы
                if cell.tables:
                    for nested_table in cell.tables:
                        row_text.append(self._read_docx_table(nested_table))
                else:
                    # Обрабатываем текст ячейки
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            cell_text += paragraph.text + " "
                    row_text.append(cell_text.strip())

            # Фильтруем пустые ячейки и объединяем строку
            filtered_row = [cell for cell in row_text if cell]
            if filtered_row:
                table_text += " | ".join(filtered_row) + '\n'

        return table_text

    def save_data(self, output_file=None):
        """Сохраняет прочитанные данные в файл"""
        if self.content is None:
            raise ValueError("Сначала необходимо прочитать файл с помощью метода read()")

        if output_file is None:
            output_file = f"output_{os.path.basename(self.name)}.txt"

        with open(output_file, 'w', encoding='utf-8') as file:
            file.write(self.content)

        return output_file

    def get_content(self):
        """Возвращает прочитанный контент"""
        return self.content

    def get_len(self):
        return len(self.content)


# Пример использования
if __name__ == "__main__":
    try:
        # Пример с текстовым файлом
        txt_reader = Readfile("golunova_n.e-sbornik_receptur_bljud_i_kulinarnykh_.docx")
        txt_reader.read()
        txt_reader.clear()
        print("Текст из TXT файла:")
        print(txt_reader.get_content()[200:300] + "...")
        print()

        # Сохранение данных
        txt_reader.save_data("saved_txt.txt")

    except Exception as e:
        print(f"Ошибка: {e}")
# import os
# import re
# import warnings
# import logging
# from pathlib import Path
# from typing import Optional, Union, List, Dict, Any
# import pandas as pd

# # Настройка подавления предупреждений
# warnings.filterwarnings("ignore")
# logging.getLogger("pypdf").setLevel(logging.ERROR)



# class ReadFile:
#     """
#     Класс для чтения текстовых файлов различных форматов.
#     С улучшенным извлечением таблиц из PDF.
#     """

#     def __init__(self, file_path: Union[str, Path]):
#         self.file_path = Path(file_path)
#         self.allowed_extensions = {'.txt', '.pdf'}
#         self.content: Optional[str] = None
#         self.tables: List[pd.DataFrame] = []
        
#     def _check_file_type(self) -> str:
#         """Проверяет тип файла и его существование."""
#         if not self.file_path.exists():
#             raise FileNotFoundError(f"Файл {self.file_path} не найден")
            
#         extension = self.file_path.suffix.lower()
#         if extension not in self.allowed_extensions:
#             raise ValueError(
#                 f"Неподдерживаемый тип файла: {extension}. "
#                 f"Разрешены: {', '.join(self.allowed_extensions)}"
#             )
            
#         return extension

#     def read(self) -> None:
#         """Читает файл в зависимости от его типа."""
#         extension = self._check_file_type()

#         try:
#             if extension == '.txt':
#                 self._read_txt()
#             elif extension == '.pdf':
#                 self._read_pdf_with_tables()
                
#         except Exception as e:
#             raise Exception(f"Ошибка при чтении файла {self.file_path}: {str(e)}")

#     def _read_txt(self) -> None:
#         """Читает текстовый файл."""
#         encodings = ['utf-8', 'cp1251', 'koi8-r', 'iso-8859-1']
        
#         for encoding in encodings:
#             try:
#                 with open(self.file_path, 'r', encoding=encoding) as file:
#                     self.content = file.read()
#                 return
#             except UnicodeDecodeError:
#                 continue
                
#         raise UnicodeDecodeError(f"Не удалось декодировать файл {self.file_path}")

#     def _read_pdf_with_tables(self) -> None:
#         """Читает PDF файл с извлечением таблиц."""
#         if not PDFPLUMBER_AVAILABLE:
#             # Если pdfplumber нет, используем обычное чтение
#             self._read_pdf_simple()
#             return
            
#         try:
#             text_parts = []
#             self.tables = []
            
#             with pdfplumber.open(self.file_path) as pdf:
#                 for page_num, page in enumerate(pdf.pages, 1):
#                     # Извлекаем обычный текст
#                     page_text = page.extract_text()
#                     if page_text:
#                         cleaned_text = self._clean_pdf_text(page_text)
#                         if cleaned_text.strip():
#                             text_parts.append(cleaned_text)
                    
#                     # Извлекаем таблицы
#                     tables = page.extract_tables()
#                     for table_num, table in enumerate(tables, 1):
#                         if table and any(any(cell for cell in row if cell) for row in table):
#                             df = self._process_table(table, page_num, table_num)
#                             self.tables.append(df)
#                             # Добавляем таблицу в текст как структурированные данные
#                             text_parts.append(f"\n--- ТАБЛИЦА {len(self.tables)} ---")
#                             text_parts.append(self._dataframe_to_text(df))
            
#             self.content = '\n\n'.join(text_parts)
                    
#         except Exception as e:
#             # Если pdfplumber не сработал, пробуем обычный метод
#             self._read_pdf_simple()

#     def _read_pdf_simple(self) -> None:
#         """Резервный метод чтения PDF без извлечения таблиц."""
#         if not PYPDF_AVAILABLE:
#             raise ImportError("pypdf/PyPDF2 не установлен")
            
#         try:
#             with warnings.catch_warnings():
#                 warnings.simplefilter("ignore")
                
#                 with open(self.file_path, 'rb') as file:
#                     pdf_reader = PdfReader(file)
#                     text_parts = []
                    
#                     for page_num, page in enumerate(pdf_reader.pages, 1):
#                         try:
#                             page_text = page.extract_text()
#                             if page_text:
#                                 cleaned_text = self._clean_pdf_text(page_text)
#                                 if cleaned_text.strip():
#                                     text_parts.append(cleaned_text)
#                         except Exception:
#                             continue
                    
#                     self.content = '\n\n'.join(text_parts)
                    
#         except Exception as e:
#             raise Exception(f"Ошибка чтения PDF: {str(e)}")

#     def _process_table(self, table: List[List[str]], page_num: int, table_num: int) -> pd.DataFrame:
#         """Обрабатывает извлеченную таблицу."""
#         # Очищаем ячейки таблицы
#         cleaned_table = []
#         for row in table:
#             cleaned_row = [self._clean_cell(cell) for cell in row if cell is not None]
#             if any(cleaned_row):  # Добавляем только непустые строки
#                 cleaned_table.append(cleaned_row)
        
#         # Создаем DataFrame
#         if cleaned_table:
#             # Используем первую строку как заголовки, если таблица имеет нормальную структуру
#             if len(cleaned_table) > 1:
#                 df = pd.DataFrame(cleaned_table[1:], columns=cleaned_table[0])
#             else:
#                 df = pd.DataFrame(cleaned_table)
#         else:
#             df = pd.DataFrame()
            
#         return df

#     def _clean_cell(self, cell: str) -> str:
#         """Очищает ячейку таблицы."""
#         if not cell:
#             return ""
#         return re.sub(r'\s+', ' ', str(cell).strip())

#     def _dataframe_to_text(self, df: pd.DataFrame) -> str:
#         """Конвертирует DataFrame в читаемый текст."""
#         if df.empty:
#             return "[Пустая таблица]"
        
#         # Используем текстовое представление DataFrame
#         return df.to_string(index=False)

#     def _clean_pdf_text(self, text: str) -> str:
#         """Очищает текст PDF от некорректных символов."""
#         if not text:
#             return ""
        
#         try:
#             if isinstance(text, bytes):
#                 text = text.decode('utf-8', errors='ignore')
            
#             cleaned = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
#             cleaned = re.sub(r'\s+', ' ', cleaned)
            
#             return cleaned.strip()
            
#         except Exception:
#             return ""

#     def get_tables(self) -> List[pd.DataFrame]:
#         """Возвращает список извлеченных таблиц."""
#         return self.tables

#     def save_tables(self, output_dir: Optional[Union[str, Path]] = None) -> List[Path]:
#         """Сохраняет таблицы в CSV файлы."""
#         if not self.tables:
#             return []
            
#         if output_dir is None:
#             output_dir = self.file_path.parent / "tables"
#         else:
#             output_dir = Path(output_dir)
            
#         output_dir.mkdir(exist_ok=True)
        
#         saved_files = []
#         for i, table in enumerate(self.tables, 1):
#             if not table.empty:
#                 output_file = output_dir / f"table_{i}.csv"
#                 table.to_csv(output_file, index=False, encoding='utf-8')
#                 saved_files.append(output_file)
                
#         return saved_files

#     def get_content(self) -> str:
#         """Возвращает прочитанный контент."""
#         if self.content is None:
#             raise ValueError("Файл еще не прочитан")
#         return self.content

#     def save_data(self, output_file: Optional[Union[str, Path]] = None) -> Path:
#         """Сохраняет прочитанные данные в файл."""
#         if self.content is None:
#             raise ValueError("Сначала необходимо прочитать файл")

#         if output_file is None:
#             output_file = self.file_path.parent / f"cleaned_{self.file_path.name}.txt"
#         else:
#             output_file = Path(output_file)

#         with open(output_file, 'w', encoding='utf-8') as file:
#             file.write(self.content)

#         return output_file


# # Пример использования
# if __name__ == "__main__":
#     try:
#         # Установите зависимости: pip install pdfplumber pandas
#         pdf_reader = ReadFile("golunova_n.e-sbornik_receptur_bljud_i_kulinarnykh_.docx")  # Замените на ваш файл
#         pdf_reader.read()
        
#         print("Успешно прочитано!")
#         print(f"Длина текста: {len(pdf_reader.get_content())}")
        
#         # Получаем таблицы
#         tables = pdf_reader.get_tables()
#         print(f"Найдено таблиц: {len(tables)}")
        
#         for i, table in enumerate(tables, 1):
#             print(f"\n--- ТАБЛИЦА {i} ---")
#             print(table)
#             print(f"Размер: {table.shape}")
        
#         # Сохраняем таблицы в CSV
#         if tables:
#             saved_files = pdf_reader.save_tables()
#             print(f"\nТаблицы сохранены в: {[str(f) for f in saved_files]}")
        
#     except Exception as e:
#         print(f"Ошибка: {e}")
#         print("\nДля извлечения таблиц установите: pip install pdfplumber pandas")